"""Document Exporter for Surya OCR.
Converts extracted page OCR results (text, tables, headings, and cropped image elements)
into styled Microsoft Word (.docx) documents.
"""

from __future__ import annotations

import io
import re
from typing import List, TYPE_CHECKING
from PIL import Image
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

if TYPE_CHECKING:
    from surya.recognition.schema import PageOCRResult


def _set_cell_background(cell, fill_hex: str):
    """Set background color of a Word table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _strip_html_tags(html_str: str) -> str:
    """Strip basic HTML tags to extract raw text."""
    if not html_str:
        return ""
    # Replace <math> tags with LaTeX text
    clean = re.sub(r"<math\b[^>]*>(.*?)</math>", r" \1 ", html_str, flags=re.DOTALL | re.IGNORECASE)
    # Remove HTML tags
    clean = re.sub(r"<[^>]+>", " ", clean)
    # Unescape common HTML entities
    clean = clean.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    return re.sub(r"\s+", " ", clean).strip()


def create_docx_from_surya_page(
    page: PageOCRResult,
    pil_image: Image.Image,
    document_title: str = "Tata Power Digitized Document",
) -> bytes:
    """Creates a high-fidelity Word (.docx) document from Surya PageOCRResult,
    including embedded cropped image blocks, tables, and styled text layout.
    """
    doc = docx.Document()

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Document Header / Banner Styling
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_run = header_p.add_run("TATA POWER DOCUMENT INTELLIGENCE HUB")
    h_run.font.name = "Arial"
    h_run.font.size = Pt(8.5)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(0, 51, 102)

    # Main Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(14)
    t_run = title_p.add_run(document_title)
    t_run.font.name = "Arial"
    t_run.font.size = Pt(22)
    t_run.font.bold = True
    t_run.font.color.rgb = RGBColor(0, 51, 102)

    img_w, img_h = pil_image.size

    for blk in page.blocks:
        x0, y0, x1, y1 = (int(c) for c in blk.bbox)
        label = blk.label or "Text"
        body_html = blk.html or ""

        # Check if this block is an Image / Visual block or skipped OCR block
        is_visual = (
            blk.skipped
            or label in ("Picture", "Figure", "Image", "Diagram", "Logo", "Stamp", "Photo")
            or not body_html.strip()
        )

        if is_visual:
            cx0 = max(0, x0 - 4)
            cy0 = max(0, y0 - 4)
            cx1 = min(img_w, x1 + 4)
            cy1 = min(img_h, y1 + 4)
            if cx1 > cx0 and cy1 > cy0:
                crop = pil_image.crop((cx0, cy0, cx1, cy1))
                buf = io.BytesIO()
                crop.save(buf, format="PNG")
                buf.seek(0)
                
                # Proportional width calculation (max 5.5 inches)
                crop_w_px = cx1 - cx0
                target_w_inches = min(5.5, max(1.2, (crop_w_px / img_w) * 6.5))
                
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                try:
                    p.add_run().add_picture(buf, width=Inches(target_w_inches))
                except Exception:
                    pass
            continue

        # Check for Tables
        if "<table>" in body_html.lower() or label in ("Table", "TableOfContents"):
            # Extract <tr> and <td>/<th>
            rows_data: List[List[str]] = []
            tr_matches = re.findall(r"<tr\b[^>]*>(.*?)</tr>", body_html, flags=re.DOTALL | re.IGNORECASE)
            for tr in tr_matches:
                cell_matches = re.findall(r"<(?:td|th)\b[^>]*>(.*?)</(?:td|th)>", tr, flags=re.DOTALL | re.IGNORECASE)
                if cell_matches:
                    rows_data.append([_strip_html_tags(c) for c in cell_matches])
            
            if rows_data:
                num_cols = max(len(r) for r in rows_data)
                word_table = doc.add_table(rows=len(rows_data), cols=num_cols)
                word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for r_idx, row_cells in enumerate(rows_data):
                    row = word_table.rows[r_idx]
                    for c_idx, cell_value in enumerate(row_cells):
                        if c_idx < num_cols:
                            cell = row.cells[c_idx]
                            cell.text = cell_value
                            if r_idx == 0:
                                _set_cell_background(cell, "003366")
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.font.name = "Arial"
                                        run.font.bold = True
                                        run.font.color.rgb = RGBColor(255, 255, 255)
                                        run.font.size = Pt(10)
                            else:
                                if r_idx % 2 == 1:
                                    _set_cell_background(cell, "F2F5F8")
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.font.name = "Arial"
                                        run.font.size = Pt(10)
                empty_p = doc.add_paragraph()
                empty_p.paragraph_format.space_after = Pt(6)
                continue

        # Headings & Text
        text_content = _strip_html_tags(body_html)
        if not text_content:
            continue

        if label == "Title":
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(text_content)
            run.font.name = "Arial"
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif label in ("SectionHeader", "Header"):
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(text_content)
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 102, 153)
        elif label in ("PageHeader", "PageFooter", "Caption"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text_content)
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            run.font.italic = True
            run.font.color.rgb = RGBColor(102, 102, 102)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text_content)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(34, 34, 34)

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_docx_from_surya_pages(
    pages: List[tuple[PageOCRResult, Image.Image]],
    document_title: str = "Tata Power Digitized Document",
) -> bytes:
    """Creates a multi-page high-fidelity Word (.docx) document from a list of (PageOCRResult, PIL.Image) tuples."""
    if not pages:
        return b""

    if len(pages) == 1:
        return create_docx_from_surya_page(pages[0][0], pages[0][1], document_title)

    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Document Header / Banner
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_run = header_p.add_run("TATA POWER DOCUMENT INTELLIGENCE HUB")
    h_run.font.name = "Arial"
    h_run.font.size = Pt(8.5)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(0, 51, 102)

    # Main Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(14)
    t_run = title_p.add_run(document_title)
    t_run.font.name = "Arial"
    t_run.font.size = Pt(22)
    t_run.font.bold = True
    t_run.font.color.rgb = RGBColor(0, 51, 102)

    for p_idx, (page, pil_image) in enumerate(pages):
        if p_idx > 0:
            doc.add_page_break()
            p_hdr = doc.add_paragraph()
            p_hdr.paragraph_format.space_before = Pt(10)
            p_hdr.paragraph_format.space_after = Pt(10)
            p_run = p_hdr.add_run(f"--- Page {p_idx + 1} ---")
            p_run.font.name = "Arial"
            p_run.font.size = Pt(10)
            p_run.font.bold = True
            p_run.font.color.rgb = RGBColor(108, 117, 125)

        for blk in page.blocks:
            x0, y0, x1, y1 = (int(c) for c in blk.bbox)
            label = blk.label or "Text"
            body_html = blk.html or ""

            is_visual = (
                blk.skipped
                or label in ("Picture", "Figure", "Image", "Diagram", "Logo", "Stamp", "Photo")
                or not body_html.strip()
            )

            if is_visual:
                pad = 4
                cx0 = max(0, x0 - pad)
                cy0 = max(0, y0 - pad)
                cx1 = min(pil_image.size[0], x1 + pad)
                cy1 = min(pil_image.size[1], y1 + pad)
                if cx1 > cx0 and cy1 > cy0:
                    crop = pil_image.crop((cx0, cy0, cx1, cy1))
                    buf = io.BytesIO()
                    crop.save(buf, format="PNG")
                    buf.seek(0)

                    img_p = doc.add_paragraph()
                    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_p.paragraph_format.space_before = Pt(8)
                    img_p.paragraph_format.space_after = Pt(8)
                    width_in = min(5.5, max(1.5, (cx1 - cx0) / 120))
                    img_p.add_run().add_picture(buf, width=Inches(width_in))
                continue

            if "<table>" in body_html.lower():
                table_data = _parse_html_table(body_html)
                if table_data:
                    n_rows = len(table_data)
                    n_cols = max(len(r) for r in table_data)
                    tbl = doc.add_table(rows=n_rows, cols=n_cols)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for r_idx, row_vals in enumerate(table_data):
                        row = tbl.rows[r_idx]
                        for c_idx, cell_value in enumerate(row_vals):
                            if c_idx < len(row.cells):
                                cell = row.cells[c_idx]
                                cell.text = cell_value
                                if r_idx == 0:
                                    _set_cell_background(cell, "003366")
                                    for p in cell.paragraphs:
                                        for run in p.runs:
                                            run.font.name = "Arial"
                                            run.font.bold = True
                                            run.font.color.rgb = RGBColor(255, 255, 255)
                                            run.font.size = Pt(10)
                                else:
                                    if r_idx % 2 == 1:
                                        _set_cell_background(cell, "F2F5F8")
                                    for p in cell.paragraphs:
                                        for run in p.runs:
                                            run.font.name = "Arial"
                                            run.font.size = Pt(10)
                    empty_p = doc.add_paragraph()
                    empty_p.paragraph_format.space_after = Pt(6)
                    continue

            text_content = _strip_html_tags(body_html)
            if not text_content:
                continue

            if label == "Title":
                h = doc.add_heading(level=1)
                h.paragraph_format.space_before = Pt(14)
                h.paragraph_format.space_after = Pt(6)
                run = h.add_run(text_content)
                run.font.name = "Arial"
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
            elif label in ("SectionHeader", "Header"):
                h = doc.add_heading(level=2)
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(4)
                run = h.add_run(text_content)
                run.font.name = "Arial"
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 85, 153)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
                run = p.add_run(text_content)
                run.font.name = "Arial"
                run.font.size = Pt(10.5)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_docx_from_markdown(
    markdown_content: str,
    document_title: str = "Tata Power Digitized Document",
) -> bytes:
    """Parses Markdown text and exports a styled Word (.docx) document."""
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_run = header_p.add_run("TATA POWER DOCUMENT INTELLIGENCE HUB")
    h_run.font.name = "Arial"
    h_run.font.size = Pt(8.5)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(0, 51, 102)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(14)
    t_run = title_p.add_run(document_title)
    t_run.font.name = "Arial"
    t_run.font.size = Pt(22)
    t_run.font.bold = True
    t_run.font.color.rgb = RGBColor(0, 51, 102)

    lines = markdown_content.splitlines()
    in_table = False
    table_lines: List[str] = []

    def flush_table(t_lines: List[str]):
        if not t_lines:
            return
        rows_data = []
        for tl in t_lines:
            if re.match(r"^\s*\|?\s*[-:]+\s*(\|\s*[-:]+\s*)*\|?\s*$", tl):
                continue
            cells = [c.strip() for c in tl.strip("|").split("|")]
            rows_data.append(cells)

        if not rows_data:
            return

        num_cols = max(len(r) for r in rows_data)
        word_table = doc.add_table(rows=len(rows_data), cols=num_cols)
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, row_cells in enumerate(rows_data):
            row = word_table.rows[r_idx]
            for c_idx, cell_value in enumerate(row_cells):
                if c_idx < num_cols:
                    cell = row.cells[c_idx]
                    cell.text = cell_value
                    if r_idx == 0:
                        _set_cell_background(cell, "003366")
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.name = "Arial"
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.size = Pt(10)
                    else:
                        if r_idx % 2 == 1:
                            _set_cell_background(cell, "F2F5F8")
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.name = "Arial"
                                run.font.size = Pt(10)

        empty_p = doc.add_paragraph()
        empty_p.paragraph_format.space_after = Pt(6)

    for line in lines:
        stripped = line.strip()

        if "|" in stripped and (stripped.startswith("|") or stripped.endswith("|")):
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table:
            in_table = False
            flush_table(table_lines)
            table_lines = []

        if not stripped:
            continue

        if stripped.startswith("# "):
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(stripped[2:].strip("# "))
            run.font.name = "Arial"
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif stripped.startswith("## "):
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(stripped[3:].strip("# "))
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 102, 153)
        elif stripped.startswith("### "):
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(stripped[4:].strip("# "))
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(51, 51, 51)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped[2:])
            run.font.name = "Arial"
            run.font.size = Pt(10.5)
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            text_part = re.sub(r"^\d+\.\s", "", stripped)
            run = p.add_run(text_part)
            run.font.name = "Arial"
            run.font.size = Pt(10.5)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            clean_text = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
            run = p.add_run(clean_text)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(34, 34, 34)

    if in_table and table_lines:
        flush_table(table_lines)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
